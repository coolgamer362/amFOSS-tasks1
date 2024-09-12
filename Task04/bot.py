import os
import requests
import csv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, MessageHandler, ConversationHandler, filters
from docx import Document
import logging

logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = "7260149709:AAGYf_b1PxW-hgylwEnN90QHDmEPzAoC5hU"
API_KEY = "AIzaSyB2CiA18mDuONkcrYw9O94NU459tJ-PgPM"

GENRE, PREVIEW, ADD, DELETE = range(4)  
reading_list = []

async def start(update: Update, context):
    await update.message.reply_text("Ahoy matey! Welcome to PagePal! Type /help to see available commands.")

async def help_command(update: Update, context):
    help_text = """
    /start - Start the bot
    /book - Get a list of books sorted by genre
    /preview - Get a preview link for a book
    /list - Manage your reading list
    /help - Show available commands
    """
    await update.message.reply_text(help_text)

async def book(update: Update, context):
    await update.message.reply_text("Please enter the genre of books you're interested in:")
    return GENRE  

async def fetch_books(update: Update, context):
    genre = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={API_KEY}')
    books = response.json().get('items', [])
    
    if not books:
        await update.message.reply_text('No books found for that genrer')
        return ConversationHandler.END

    with open('books.csv', mode='w', newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Title", "Author", "Description", "Year", "Language", "Preview Link"])
        
        for book in books:
            info = book.get('volumeInfo', {})
            title = info.get('title', 'N/A')
            authors = ", ".join(info.get('authors', []))
            description = info.get('description', 'N/A')
            year = info.get('publishedDate', 'N/A')
            language = info.get('language', 'N/A')
            preview = info.get('previewLink', 'N/A')
            writer.writerow([title, authors, description, year, language, preview])
    
    await update.message.reply_document(document=open('books.csv', 'rb'))
    return ConversationHandler.END  

async def preview(update: Update, context):
    await update.message.reply_text("Please enter the name of the book to previeew:")
    return PREVIEW  

async def fetch_preview(update: Update, context):
    book_name = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={API_KEY}')
    books = response.json().get('items', [])
    
    if books:
        preview_link = books[0].get('volumeInfo', {}).get('previewLink', 'No preview available')
        await update.message.reply_text(f'Preview link: {preview_link}')
    else:
        await update.message.reply_text('No books found with that name.')
    return ConversationHandler.END  

async def list_books(update: Update, context):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add_book')],
        [InlineKeyboardButton("Delete a book", callback_data='delete_book')],
        [InlineKeyboardButton("View Reading List", callback_data='view_list')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Manage your reading list:', reply_markup=reply_markup)

async def handle_buttons(update: Update, context):
    query = update.callback_query
    await query.answer()

    if query.data == 'add_book':
        await query.message.reply_text("Enter the name of the book to add:")
        return ADD  
    elif query.data == 'delete_book':
        await query.message.reply_text("Enter the name of the book to delette:")
        return DELETE  
    elif query.data == 'view_list':
        doc = Document()
        doc.add_heading('Reading List', 0)
        for book in reading_list:
            doc.add_paragraph(f"Title: {book['title']}")
            doc.add_paragraph(f"Preview: {book['preview']}")
            doc.add_paragraph('---')
        doc.save('reading_list.docx')
        await query.message.reply_document(document=open('reading_list.docx', 'rb'))
    return ConversationHandler.END

async def add_book(update: Update, context):
    book_name = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={API_KEY}')
    books = response.json().get('items', [])
    
    if books:
        book_info = books[0].get('volumeInfo', {})
        reading_list.append({
            'title': book_info.get('title', 'N/A'),
            'preview': book_info.get('previewLink', 'N/A')
        })
        await update.message.reply_text(f'Book "{book_name}" added to the reading list.')
    else:
        await update.message.reply_text('No books found with that name.')
    return ConversationHandler.END  

async def delete_book(update: Update, context):
    book_name = update.message.text
    global reading_list
    reading_list = [book for book in reading_list if book['title'].lower() != book_name.lower()]
    await update.message.reply_text(f'Book "{book_name}" deletted from the reading list.')
    return ConversationHandler.END  

if __name__ == "__main__":
    app = ApplicationBuilder().token(BOT_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler("book", book),
            CommandHandler("preview", preview),
            CallbackQueryHandler(handle_buttons),
        ],
        states={
            GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, fetch_books)],
            PREVIEW: [MessageHandler(filters.TEXT & ~filters.COMMAND, fetch_preview)],
            ADD: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_book)],
            DELETE: [MessageHandler(filters.TEXT & ~filters.COMMAND, delete_book)],
        },
        fallbacks=[CommandHandler("help", help_command)]
    )

    app.add_handler(CommandHandler("start", start))
    app.add_handler(CommandHandler("help", help_command))
    app.add_handler(CommandHandler("list", list_books))
    app.add_handler(conv_handler)

    app.run_polling()

